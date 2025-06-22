from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import tempfile
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import openai

app = Flask(__name__)
CORS(app)

# OpenAI 1.0.0+ үшін дұрыс инициализация
from openai import OpenAI
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))


def build_prompt(topic, subject, grade, language_level, bloom_level):
    return f"""
You are a CLIL lesson planner for school teachers in Kazakhstan.

Create a full CLIL-based lesson plan for the subject "{subject}" on the topic "{topic}" for Grade {grade} students.
The learners' English level is {language_level}, and the cognitive focus should be based on Bloom's level: {bloom_level}.

The lesson plan must be structured as a 2-column table:
Column 1 — Section Name (e.g., Objectives, Assessment, Language Focus, etc.)
Column 2 — Content for each section (described in clear paragraphs)

Sections to include:
- Title of lesson
- Grade level
- Learning objectives (content + language)
- Assessment criteria
- Subject vocabulary (EN + KZ)
- Bloom’s level
- 4Cs focus (Content, Communication, Cognition, Culture)
- Pre-knowledge
- Lesson stages (Beginning – Middle – End)
- Differentiation
- Values
- ICT used
- Resources (include specific useful websites, tools, programs, and platforms with names and links that teachers can visit directly to use in class)

Output format must be a clean textual table with each row representing a section, using ":" between section name and content.
"""


@app.route("/generate_lessonplan", methods=["POST"])
def generate_lessonplan():
    data = request.json
    topic = data.get("topic", "")
    subject = data.get("subject", "Informatics")  # Үнемі "Informatics"
    grade = data.get("grade", "")
    language_level = data.get("language_level", "")
    bloom_level = data.get("bloom_level", "")

    prompt = build_prompt(topic, subject, grade, language_level, bloom_level)

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a CLIL methodology expert generating professional lesson plans."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7
    )

    result = response.choices[0].message.content
    return jsonify({"lesson_plan": result})


@app.route("/download_lessonplan_docx", methods=["POST"])
def download_lessonplan_docx():
    data = request.json
    content = data.get("lesson_plan", "")
    doc = Document()
    
    # Ортақ стиль орнату (Қаріп, өлшем)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # Қазақша үшін қаріпті орнату (қажет болса)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.add_heading("CLIL Сабақ жоспары", level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    for line in content.strip().split('\n'):
        if ':' in line:
            col1, col2 = line.split(':', 1)
            row_cells = table.add_row().cells
            row_cells[0].text = col1.strip()
            row_cells[1].text = col2.strip()
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return send_file(temp_file.name, as_attachment=True, download_name="lesson_plan.docx")


@app.route("/download_lessonplan_pdf", methods=["POST"])
def download_lessonplan_pdf():
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors

    data = request.json
    content = data.get("lesson_plan", "")

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")

    # Параққа тақырыпты қою
    doc = SimpleDocTemplate(temp_file.name, pagesize=letter,
                            rightMargin=20, leftMargin=20,
                            topMargin=20, bottomMargin=20)
    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]
    normal_style.fontName = "Helvetica"
    normal_style.fontSize = 11
    normal_style.leading = 14

    elements = []

    # Тақырып
    elements.append(Paragraph("CLIL Сабақ жоспары", styles["Title"]))
    elements.append(Spacer(1, 12))

    # Кестеге мәтінді бөлшектеу
    data_table = []
    for line in content.strip().split('\n'):
        if ':' in line:
            col1, col2 = line.split(':', 1)
            data_table.append([col1.strip(), col2.strip()])
        else:
            # Егер жолда ':' болмаса, жайша бөлек параграф қосуға болады
            data_table.append([line.strip(), ""])

    table_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ('BOX', (0, 0), (-1, -1), 0.25, colors.grey),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ])

    table = Table(data_table, colWidths=[100*mm, 80*mm])
    table.setStyle(table_style)

    elements.append(table)

    doc.build(elements)

    return send_file(temp_file.name, as_attachment=True, download_name="lesson_plan.pdf")


if __name__ == "__main__":
    app.run(debug=True)
